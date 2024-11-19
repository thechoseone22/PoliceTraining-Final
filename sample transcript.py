from docx import Document
import pandas as pd
df = pd.read_csv("cleaned_dataset.csv")

# pick an example transcript
example_transcript = df["Bodycam Transcript"].iloc[0]  # Replace with the row index you want

# create word document
doc = Document()
doc.add_heading("AI-Generated Example Transcript", level=1)
doc.add_paragraph(example_transcript)

# save
doc.save("Example_Transcript.docx")
print("you did it.")
